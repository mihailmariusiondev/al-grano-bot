from docx import Document
import PyPDF2
import io
import logging
from typing import Optional
from bot.services.openai_service import openai_service

logger = logging.getLogger(__name__)

SUPPORTED_DOCUMENT_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


async def document_handler(message, context) -> Optional[str]:
    """Handle document messages"""
    try:
        document = message.document
        if document.mime_type not in SUPPORTED_DOCUMENT_TYPES:
            error_msg = f"Tipo de documento no soportado: {document.mime_type}"
            logger.warning(error_msg)
            await message.reply_text(
                "Lo siento, este tipo de documento no está soportado. Por favor, envía un PDF, DOCX o TXT."
            )
            return None

        logger.info(
            f"Procesando documento: {document.file_name} ({document.mime_type})"
        )

        # Download document
        try:
            file = await context.bot.get_file(document.file_id)
            file_bytes = await file.download_as_bytearray()
            logger.info(f"Documento descargado, tamaño: {len(file_bytes)} bytes")
        except Exception as e:
            error_msg = f"Error descargando el documento: {str(e)}"
            logger.error(error_msg, exc_info=True)
            await message.reply_text(
                "Lo siento, hubo un problema al descargar el documento. Por favor, inténtalo de nuevo."
            )
            return None

        # Extract text based on file type
        try:
            if document.mime_type == "application/pdf":
                text_content = extract_text_from_pdf(file_bytes)
            elif (
                document.mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                text_content = extract_text_from_docx(file_bytes)
            else:  # text/plain
                text_content = file_bytes.decode("utf-8")

            if not text_content or not text_content.strip():
                logger.error("No se pudo extraer texto del documento")
                await message.reply_text(
                    "No pude extraer texto de este documento. ¿Podrías verificar que no esté vacío o dañado?"
                )
                return None

            logger.info(
                f"Texto extraído exitosamente, longitud: {len(text_content)} caracteres"
            )
            return text_content

        except Exception as e:
            error_msg = f"Error extrayendo texto del documento: {str(e)}"
            logger.error(error_msg, exc_info=True)
            await message.reply_text(
                "Hubo un problema al procesar el documento. ¿Podrías intentar con otro archivo?"
            )
            return None

    except Exception as e:
        logger.error(f"Error en document handler: {str(e)}", exc_info=True)
        await message.reply_text(
            "Ocurrió un error inesperado al procesar el documento. Por favor, inténtalo de nuevo."
        )
        return None


def extract_text_from_pdf(file_bytes: bytearray) -> str:
    """Extract text content from PDF file"""
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_file)
        text_content = []

        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                logger.debug(
                    f"Página {page_num + 1} procesada: {len(page_text)} caracteres"
                )
            except Exception as e:
                logger.warning(f"Error al procesar página {page_num + 1}: {e}")
                continue

        full_text = "\n\n".join(text_content)
        if not full_text.strip():
            raise ValueError("No se pudo extraer texto del PDF")

        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}", exc_info=True)
        raise


def extract_text_from_docx(file_bytes: bytearray) -> str:
    """Extract text content from DOCX file"""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text_content = []

        # Procesar párrafos
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                logger.debug(
                    f"Párrafo {i + 1} procesado: {len(paragraph.text)} caracteres"
                )
            except Exception as e:
                logger.warning(f"Error al procesar párrafo {i + 1}: {e}")
                continue

        # Procesar tablas
        for i, table in enumerate(doc.tables):
            try:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_content.append(row_text)
                logger.debug(f"Tabla {i + 1} procesada")
            except Exception as e:
                logger.warning(f"Error al procesar tabla {i + 1}: {e}")
                continue

        full_text = "\n\n".join(text_content)
        if not full_text.strip():
            raise ValueError("No se pudo extraer texto del documento DOCX")

        logger.info(
            f"DOCX procesado exitosamente. Texto extraído: {len(full_text)} caracteres"
        )
        return full_text

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
        raise ValueError("No se pudo extraer contenido del documento DOCX")
